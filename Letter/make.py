from docx import Document
import json
class Letter:
    def __init__(self, TO=None, isAssistantComm=None, WardNo=None, Assessee=None, SUBJECT=None, RefNo=None, BODY=None):
        self.properties = {
            "TO" : "",
            "isAssistantComm" : False,
            "WardNo" : "",
            "Assessee" : "",
            "Address" : "",
            "PAN" : "",
            "MobNo" : "",
            "SUBJECT" : "",
            "RefNo" : "",
            "BODY" : ""
        }
        if TO:
            self.properties['TO'] = TO
        if isAssistantComm:
            self.properties['isAssistantComm'] = isAssistantComm
        if WardNo:
            self.properties['WardNo'] = WardNo
        if Assessee:
            self.properties['Assessee'] = Assessee
            self.properties['Address'] = self.get_address_from_assessee(Assessee)
            self.properties['PAN'] = self.get_pan_from_assess(Assessee)
            self.properties['MobNo'] = self.get_mobno_from_assess(Assessee)

        if SUBJECT:
            self.properties['SUBJECT'] = SUBJECT
        if RefNo:
            self.properties['RefNo'] = RefNo
        if BODY:
            self.properties['BODY'] = BODY

    def create_object_from_json(self, obj):
        return Letter(obj.get('TO'), obj.get('isAssistantComm'), obj.get('WardNo'), obj.get('Assess'), obj.get('SUBJECT'), obj.get('RefNo'), obj.get('BODY'))
    
    def get_address_from_assessee(self, Assessee):
        return "This is a demo Address"
    
    def get_mobno_from_assess(self, Assessee):
        return "+91 90090 90909"

    def get_pan_from_assess(self, Assessee):
        return "ABCD89EIJK"

    def save_to_docs(self, path):
        # letter_string = ""
        # with open('Letter/demo_letter.txt', 'r') as fp:
        #     letter_string = fp.read()
        # WardString = ""
        # AssistantComm = ""
        # if(self.properties['WardNo']):
        #     CityCondition = "Akola."
        #     WardString = "Ward - " + str(self.properties['WardNo']) + ","
        # if(self.properties['isAssistantComm']):
        #     CityCondition = "Akola Circle, Akola."
        #     WardString = ""
        #     AssistantComm = "Assistant Commissioner of Income Tax,"
        # letter_string = letter_string.format(
        #     TO=self.properties['TO'],
        #     WardNo=WardString,
        #     AssistantComm=AssistantComm,
        #     CityCondition=CityCondition,
        #     assessee=self.properties.get('Assessee'),
        #     address=self.properties.get('Address'),
        #     MobNo=self.properties.get('MobNo'),
        #     PAN=self.properties.get('PAN'),
        #     SUBJECT=self.properties.get('SUBJECT'),
        #     RefNo=self.properties.get('RefNo'),
        #     BODY=self.properties.get('BODY')
        #     )
        document = Document()
        para1 = document.add_paragraph('To,\n' + self.properties['TO'])
        if(self.properties['isAssistantComm']):
            para1.add_run('Assistant Commissioner of Income Tax,\n')
            para1.add_run('Akola Circle, Akola.\n\n')
        else:
            para1.add_run("Ward - " + str(self.properties['WardNo']) + ",\n")
            para1.add_run("Akola.\n\n")
        
        para2 = document.add_paragraph("Assessee\t:\t" + self.properties.get('Assessee') + "\n")
        para2.add_run("\t\t\t" + self.properties.get('Address') + "\n")
        para2.add_run("\t\t\t" + 'Cell No. : ' + self.properties.get('MobNo') + "\n\n")

        para3 = document.add_paragraph("PAN\t:\t" + self.properties.get('PAN') + "\n")
        para4 = document.add_paragraph("Subject\t:\t" + self.properties.get('SUBJECT') + "\n")
        para5 = document.add_paragraph("Reference No.\t:\t" + self.properties.get('RefNo') + "\n\n")
        para6 = document.add_paragraph(self.properties.get('BODY') + "\n\n")
        para7 = document.add_paragraph("Yours Faithfully,\n\n---------------------")

        # document.add_paragraph(letter_string)
        document.save(path)
    
def get_choices_kvp():
    fp = open("C:\\Users\\kismatit\\OneDrive\\Desktop\\LetterGen\\data.json", "r")
    raw_data = fp.read()
    data = json.loads(raw_data)
    kvps = []
    for d in data:
        k = d['id']
        v = d['name']
        kvps.append((k, v))
    return kvps

# letter = Letter("The Income Tax Officer", False, 6, "SOME ASSESSEE", "This letter is regarding Demo", "909", "This is a demo test letter.\nSome XYZ, Testing Currently")
# letter.save_to_docs('demo.docx')